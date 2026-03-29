"""创建一个充满 AI 味的测试 Word 文档"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path


def create_ai_flavored_doc(output_path: str):
    doc = Document()

    # 标题
    title = doc.add_heading("人工智能技术在现代教育中的应用与展望", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 摘要
    doc.add_heading("摘要", level=2)
    doc.add_paragraph(
        "在当今数字化时代，人工智能技术正在深刻地改变着教育领域的方方面面。"
        "值得注意的是，随着深度学习和自然语言处理技术的不断发展，"
        "智能教育系统已经从简单的辅助工具演变为能够提供个性化学习体验的综合平台。"
        "本文旨在全面探讨人工智能技术在现代教育中的多维度应用，"
        "并对其未来发展趋势进行深入分析。"
    )

    # 正文
    doc.add_heading("一、引言", level=2)
    doc.add_paragraph(
        "首先，我们需要认识到，教育是人类社会发展的基石。"
        "在这个背景下，人工智能技术的引入无疑为教育带来了前所未有的变革。"
        "此外，随着互联网技术的普及和移动设备的广泛使用，"
        "在线教育已经成为一种不可忽视的教育形式。"
        "综上所述，研究人工智能在教育领域的应用具有重要的理论和实践意义。"
    )

    doc.add_heading("二、人工智能在教育中的主要应用", level=2)
    doc.add_paragraph(
        "对于人工智能在教育领域的应用而言，可以从以下几个方面进行分析。"
        "首先，智能辅导系统能够根据学生的学习进度和知识掌握情况，"
        "提供个性化的学习路径推荐。"
        "其次，自然语言处理技术使得智能问答系统成为可能，"
        "学生可以随时向系统提问并获得即时反馈。"
        "最后，计算机视觉技术在在线考试监控和作业批改方面也发挥着重要作用。"
    )

    doc.add_heading("三、面临的挑战与问题", level=2)
    doc.add_paragraph(
        "尽管人工智能在教育领域展现出巨大的潜力，"
        "但我们也必须清醒地认识到其面临的诸多挑战。"
        "值得注意的是，数据隐私和安全问题是当前亟待解决的关键问题之一。"
        "此外，技术的公平性和可及性也是我们需要重点关注的方面。"
        "在数字鸿沟依然存在的背景下，"
        "如何确保每个学生都能平等地享受到人工智能带来的教育红利，"
        "是一个值得深入思考的问题。"
    )

    doc.add_heading("四、未来展望", level=2)
    doc.add_paragraph(
        "展望未来，人工智能技术在教育领域的应用前景是令人振奋的。"
        "一方面，随着技术的不断进步，智能教育系统将变得更加智能和人性化。"
        "另一方面，教育工作者需要不断提升自身的数字素养，"
        "以更好地适应和利用这些新技术。"
        "总而言之，人工智能与教育的深度融合将为构建更加公平、高效的教育体系"
        "提供强有力的技术支撑。"
    )

    doc.add_heading("结论", level=2)
    doc.add_paragraph(
        "综上所述，人工智能技术在现代教育中的应用已经取得了显著的成果，"
        "但仍然面临着诸多挑战。"
        "通过本文的分析可以看出，只有在充分认识到技术局限性的基础上，"
        "积极探索创新性的应用模式，才能真正实现人工智能与教育的深度融合。"
        "未来，我们期待看到更多创新性的教育技术应用，"
        "为每一位学习者提供更加优质的教育体验。"
    )

    # 统一格式（典型 AI 生成的完美均匀排版）
    for para in doc.paragraphs:
        para.paragraph_format.space_after = Pt(12)
        para.paragraph_format.line_spacing = 1.5
        for run in para.runs:
            run.font.size = Pt(12)
            run.font.name = "宋体"

    doc.save(output_path)
    print(f"测试文档已创建: {output_path}")
    print(f"  段落数: {len(doc.paragraphs)}")
    total_chars = sum(len(p.text) for p in doc.paragraphs)
    print(f"  总字数: {total_chars}")

    # 检查 AI 味指标
    text = "\n".join(p.text for p in doc.paragraphs)
    ai_indicators = {
        "此外": text.count("此外"),
        "值得注意的是": text.count("值得注意的是"),
        "综上所述": text.count("综上所述"),
        "首先...其次...最后": 1 if "首先" in text and "其次" in text and "最后" in text else 0,
        "在...背景下": text.count("背景下"),
        "对于...而言": text.count("而言"),
    }
    print(f"  AI 味指标: {ai_indicators}")


if __name__ == "__main__":
    output = Path("backend/tests/fixtures/ai_flavored_doc.docx")
    output.parent.mkdir(parents=True, exist_ok=True)
    create_ai_flavored_doc(str(output))
