"""DocxProcessor 单元测试 + 处理器工厂测试"""

import os
import tempfile
from pathlib import Path

import pytest
from docx import Document

from backend.processors import get_processor
from backend.processors.docx_processor import DocxProcessor

FIXTURES = Path(__file__).parent / "fixtures"


def _create_simple_doc(path: str, paragraphs: list[str]):
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(path)


def test_parse():
    processor = DocxProcessor()
    result = processor.parse(str(FIXTURES / "ai_flavored_doc.docx"))
    assert result["format"] == "docx"
    assert result["paragraph_count"] > 0
    assert len(result["paragraphs"]) > 0
    assert "text" in result["paragraphs"][0]


def test_extract_text():
    processor = DocxProcessor()
    text = processor.extract_text(str(FIXTURES / "ai_flavored_doc.docx"))
    assert len(text) > 100
    assert "人工智能" in text


def test_get_stats():
    processor = DocxProcessor()
    stats = processor.get_stats(str(FIXTURES / "ai_flavored_doc.docx"))
    assert stats["paragraph_count"] > 0
    assert stats["total_chars"] > 0


def test_replace_text():
    processor = DocxProcessor()
    with tempfile.TemporaryDirectory() as tmpdir:
        src = os.path.join(tmpdir, "src.docx")
        out = os.path.join(tmpdir, "out.docx")
        _create_simple_doc(src, ["第一段原文", "第二段原文", "第三段原文"])

        processor.replace_text(src, {"0": "第一段修改后", "2": "第三段修改后"}, out)

        doc = Document(out)
        texts = [p.text for p in doc.paragraphs]
        assert texts[0] == "第一段修改后"
        assert texts[1] == "第二段原文"
        assert texts[2] == "第三段修改后"


def test_apply_format_changes():
    processor = DocxProcessor()
    with tempfile.TemporaryDirectory() as tmpdir:
        src = os.path.join(tmpdir, "src.docx")
        out = os.path.join(tmpdir, "out.docx")
        _create_simple_doc(src, ["段落一", "段落二", "段落三"])

        changes = [
            {"target": "paragraph_spacing", "scope": "all", "value_pt": 14},
            {"target": "line_spacing", "scope": [0, 2], "value": 1.25},
        ]
        processor.apply_format_changes(src, changes, out)

        assert os.path.exists(out)
        doc = Document(out)
        assert len(doc.paragraphs) == 3


def test_cjk_font():
    processor = DocxProcessor()
    with tempfile.TemporaryDirectory() as tmpdir:
        src = os.path.join(tmpdir, "src.docx")
        out = os.path.join(tmpdir, "out.docx")
        _create_simple_doc(src, ["中文测试段落"])

        changes = [{"target": "font_name", "scope": "all", "value": "微软雅黑"}]
        processor.apply_format_changes(src, changes, out)

        doc = Document(out)
        run = doc.paragraphs[0].runs[0]
        from docx.oxml.ns import qn

        r_fonts = run._element.rPr.find(qn("w:rFonts"))
        assert r_fonts is not None
        assert r_fonts.get(qn("w:eastAsia")) == "微软雅黑"


def test_get_processor_factory():
    from backend.processors.docx_processor import DocxProcessor
    from backend.processors.pdf_processor import PdfProcessor
    from backend.processors.pptx_processor import PptxProcessor
    from backend.processors.xlsx_processor import XlsxProcessor

    assert isinstance(get_processor("test.docx"), DocxProcessor)
    assert isinstance(get_processor("test.pptx"), PptxProcessor)
    assert isinstance(get_processor("test.xlsx"), XlsxProcessor)
    assert isinstance(get_processor("test.pdf"), PdfProcessor)
    assert isinstance(get_processor("/path/to/FILE.DOCX"), DocxProcessor)

    with pytest.raises(ValueError, match="Unsupported format"):
        get_processor("test.txt")
