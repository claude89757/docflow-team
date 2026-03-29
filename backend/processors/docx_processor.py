from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


class DocxProcessor:
    """Word (.docx) 文档处理器"""

    def parse(self, file_path: str) -> dict:
        """解析 docx 文档结构"""
        doc = Document(file_path)
        paragraphs = []
        for i, para in enumerate(doc.paragraphs):
            paragraphs.append(
                {
                    "index": i,
                    "text": para.text,
                    "style": para.style.name if para.style else None,
                    "alignment": str(para.alignment) if para.alignment else None,
                    "runs": [
                        {
                            "text": run.text,
                            "bold": run.bold,
                            "italic": run.italic,
                            "font_name": run.font.name,
                            "font_size": str(run.font.size) if run.font.size else None,
                        }
                        for run in para.runs
                    ],
                }
            )

        tables = []
        for i, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                rows.append(cells)
            tables.append({"index": i, "rows": rows})

        return {
            "file_path": file_path,
            "format": "docx",
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
            "paragraphs": paragraphs,
            "tables": tables,
        }

    def extract_text(self, file_path: str) -> str:
        """提取纯文本"""
        doc = Document(file_path)
        texts = []
        for para in doc.paragraphs:
            if para.text.strip():
                texts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        texts.append(cell.text)
        return "\n\n".join(texts)

    def get_stats(self, file_path: str) -> dict:
        """文档统计"""
        doc = Document(file_path)
        total_chars = 0
        paragraph_count = 0
        for para in doc.paragraphs:
            if para.text.strip():
                paragraph_count += 1
                total_chars += len(para.text)
        return {
            "paragraph_count": paragraph_count,
            "table_count": len(doc.tables),
            "total_chars": total_chars,
            "image_count": len(doc.inline_shapes),
        }

    def replace_paragraph_text(self, file_path: str, replacements: dict[int, str], output_path: str) -> str:
        """替换指定段落的文本，保留格式"""
        doc = Document(file_path)
        for idx_str, new_text in replacements.items():
            idx = int(idx_str)
            if 0 <= idx < len(doc.paragraphs):
                para = doc.paragraphs[idx]
                # 保留第一个 run 的格式，清除其他 run
                if para.runs:
                    first_run = para.runs[0]
                    # 保存格式
                    bold = first_run.bold
                    italic = first_run.italic
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    # 清除所有 run
                    for run in para.runs:
                        run.text = ""
                    # 设置新文本到第一个 run
                    first_run.text = new_text
                    first_run.bold = bold
                    first_run.italic = italic
                    if font_name:
                        first_run.font.name = font_name
                    if font_size:
                        first_run.font.size = font_size
                    # 处理 CJK 字体
                    self._set_cjk_font(first_run, font_name)
                else:
                    para.text = new_text

        doc.save(output_path)
        return output_path

    def apply_format_changes(self, file_path: str, changes: list[dict], output_path: str) -> str:
        """应用格式修改指令 (JSON → 确定性执行)

        changes 格式:
        [
            {"target": "paragraph_spacing", "scope": "all", "value_pt": 13.5},
            {"target": "font_size", "scope": [3, 5, 7], "value_pt": 12},
            {"target": "line_spacing", "scope": "all", "value": 1.15},
        ]
        """
        doc = Document(file_path)

        for change in changes:
            target = change.get("target")
            scope = change.get("scope", "all")

            if target == "paragraph_spacing":
                value = Pt(change["value_pt"])
                paras = self._resolve_scope(doc.paragraphs, scope)
                for para in paras:
                    fmt = para.paragraph_format
                    fmt.space_after = value

            elif target == "line_spacing":
                value = change["value"]
                paras = self._resolve_scope(doc.paragraphs, scope)
                for para in paras:
                    para.paragraph_format.line_spacing = value

            elif target == "font_size":
                value = Pt(change["value_pt"])
                paras = self._resolve_scope(doc.paragraphs, scope)
                for para in paras:
                    for run in para.runs:
                        run.font.size = value

            elif target == "font_name":
                name = change["value"]
                paras = self._resolve_scope(doc.paragraphs, scope)
                for para in paras:
                    for run in para.runs:
                        run.font.name = name
                        self._set_cjk_font(run, name)

        doc.save(output_path)
        return output_path

    def _resolve_scope(self, paragraphs, scope):
        """解析 scope: "all" 或索引列表"""
        if scope == "all":
            return list(paragraphs)
        if isinstance(scope, list):
            return [paragraphs[i] for i in scope if 0 <= i < len(paragraphs)]
        return []

    def _set_cjk_font(self, run, font_name: str | None):
        """设置 CJK 字体 (python-docx 高级 API 不支持 eastAsia 字体)"""
        if not font_name:
            return
        rpr = run._element.get_or_add_rPr()
        r_fonts = rpr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = run._element.makeelement(qn("w:rFonts"), {})
            rpr.insert(0, r_fonts)
        r_fonts.set(qn("w:eastAsia"), font_name)
