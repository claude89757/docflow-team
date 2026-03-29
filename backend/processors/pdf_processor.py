import fitz  # PyMuPDF
from docx import Document


class PdfProcessor:
    """PDF 文档处理器

    解析和提取文本使用 PyMuPDF。
    内容修改通过转换为 docx 后委托 DocxProcessor 处理。
    """

    def parse(self, file_path: str) -> dict:
        """解析 PDF 文档结构"""
        doc = fitz.open(file_path)
        pages = []
        for i, page in enumerate(doc):
            blocks: list[dict] = []
            for block in page.get_text("dict")["blocks"]:
                if block["type"] != 0:  # 只处理文本块
                    continue
                text_parts = []
                font_sizes = []
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        text_parts.append(span["text"])
                        font_sizes.append(span["size"])
                text = " ".join(text_parts).strip()
                if text:
                    blocks.append(
                        {
                            "index": len(blocks),
                            "text": text,
                            "bbox": list(block["bbox"]),
                            "avg_font_size": round(sum(font_sizes) / len(font_sizes), 1) if font_sizes else 0,
                        }
                    )
            pages.append({"index": i, "blocks": blocks})
        doc.close()

        return {
            "file_path": file_path,
            "format": "pdf",
            "page_count": len(pages),
            "pages": pages,
        }

    def extract_text(self, file_path: str) -> str:
        """提取纯文本"""
        doc = fitz.open(file_path)
        try:
            texts = []
            for page in doc:
                text = page.get_text().strip()
                if text:
                    texts.append(text)
            return "\n\n".join(texts)
        finally:
            doc.close()

    def get_stats(self, file_path: str) -> dict:
        """文档统计"""
        doc = fitz.open(file_path)
        try:
            page_count = len(doc)
            block_count = 0
            total_chars = 0
            for page in doc:
                for block in page.get_text("blocks"):
                    if block[6] == 0:  # text block
                        block_count += 1
                        total_chars += len(block[4])
            return {
                "page_count": page_count,
                "block_count": block_count,
                "total_chars": total_chars,
            }
        finally:
            doc.close()

    def to_docx(self, file_path: str, output_path: str) -> str:
        """将 PDF 文本内容转换为 docx

        根据字号推断标题/正文层级。
        """
        pdf_doc = fitz.open(file_path)
        docx_doc = Document()

        # 收集所有文本块及其字号
        all_blocks = []
        for page in pdf_doc:
            for block in page.get_text("dict")["blocks"]:
                if block["type"] != 0:
                    continue
                lines_text = []
                font_sizes = []
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        lines_text.append(span["text"])
                        font_sizes.append(span["size"])
                text = " ".join(lines_text).strip()
                avg_size = sum(font_sizes) / len(font_sizes) if font_sizes else 12
                if text:
                    all_blocks.append((text, avg_size))

        if not all_blocks:
            pdf_doc.close()
            docx_doc.save(output_path)
            return output_path

        # 计算字号分布，推断层级：出现频率最高的字号作为正文
        body_size = max(set(s for _, s in all_blocks), key=lambda s: sum(1 for _, fs in all_blocks if fs == s))

        for text, size in all_blocks:
            if size > body_size * 1.4:
                docx_doc.add_heading(text, level=1)
            elif size > body_size * 1.15:
                docx_doc.add_heading(text, level=2)
            else:
                docx_doc.add_paragraph(text)

        pdf_doc.close()
        docx_doc.save(output_path)
        return output_path

    def replace_text(self, file_path: str, replacements: dict, output_path: str) -> str:
        """通过转 docx 实现文本替换

        先将 PDF 转为 docx，再用 DocxProcessor 替换。
        output_path 扩展名会被强制改为 .docx。
        """
        from pathlib import Path

        from backend.processors.docx_processor import DocxProcessor

        # 确保输出为 docx
        docx_output = str(Path(output_path).with_suffix(".docx"))
        temp_docx = docx_output + ".tmp.docx"

        self.to_docx(file_path, temp_docx)
        docx_proc = DocxProcessor()
        result = docx_proc.replace_text(temp_docx, replacements, docx_output)

        Path(temp_docx).unlink(missing_ok=True)
        return result

    def apply_format_changes(self, file_path: str, changes: list[dict], output_path: str) -> str:
        """通过转 docx 实现格式修改

        先将 PDF 转为 docx，再用 DocxProcessor 应用格式。
        output_path 扩展名会被强制改为 .docx。
        """
        from pathlib import Path

        from backend.processors.docx_processor import DocxProcessor

        docx_output = str(Path(output_path).with_suffix(".docx"))
        temp_docx = docx_output + ".tmp.docx"

        self.to_docx(file_path, temp_docx)
        docx_proc = DocxProcessor()
        result = docx_proc.apply_format_changes(temp_docx, changes, docx_output)

        Path(temp_docx).unlink(missing_ok=True)
        return result
